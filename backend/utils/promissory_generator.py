import os
import tempfile
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

def num_to_words(number):
    try:
        num = int(float(str(number).replace(',', '')))
    except (ValueError, TypeError):
        return ""
    
    if num == 0:
        return "Zero"
        
    def helper(n):
        units = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine", "Ten", 
                 "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen", "Seventeen", "Eighteen", "Nineteen"]
        tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        
        if n < 20:
            return units[n]
        elif n < 100:
            return tens[n // 10] + (" " + units[n % 10] if n % 10 != 0 else "")
        elif n < 1000:
            return units[n // 100] + " Hundred" + (" and " + helper(n % 100) if n % 100 != 0 else "")
        elif n < 100000:
            return helper(n // 1000) + " Thousand" + (" " + helper(n % 1000) if n % 1000 != 0 else "")
        elif n < 10000000:
            return helper(n // 100000) + " Lakh" + (" " + helper(n % 100000) if n % 100000 != 0 else "")
        else:
            return helper(n // 10000000) + " Crore" + (" " + helper(n % 10000000) if n % 10000000 != 0 else "")
            
    return helper(num) + " Rupees Only"

def format_date_suffix(date_str):
    if not date_str:
        return "this _____ day of ________, 20__"
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
    except Exception:
        return date_str
    day = dt.day
    if 11 <= day <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return dt.strftime(f"{day}{suffix} day of %B, %Y")

def calculate_installment(amount, rate_pct, no_of_periods, period_type):
    try:
        p = float(str(amount).replace(',', ''))
        r_annual = float(rate_pct)
        n = int(no_of_periods)
    except (ValueError, TypeError):
        return "0.00"
        
    periods_per_year = {
        'daily': 365,
        'weekly': 52,
        'biweekly': 26,
        'bimonthly': 6,
        'monthly': 12
    }.get(str(period_type).lower(), 12)
    
    r = r_annual / 100.0 / periods_per_year
    
    if r == 0:
        return f"{round(p / n, 2):,.2f}"
    
    try:
        # Standard reducing EMI formula
        emi = p * r * ((1 + r) ** n) / (((1 + r) ** n) - 1)
        return f"{round(emi, 2):,.2f}"
    except Exception:
        # Simple flat fallback
        total_interest = p * (r_annual / 100.0) * (n / periods_per_year)
        return f"{round((p + total_interest) / n, 2):,.2f}"

def format_indian_currency(amount_str):
    try:
        val = int(float(str(amount_str).replace(',', '')))
        val_str = str(val)
        if len(val_str) <= 3:
            return val_str
        last_three = val_str[-3:]
        remaining = val_str[:-3]
        reversed_rem = remaining[::-1]
        grouped = ",".join(reversed_rem[i:i+2] for i in range(0, len(reversed_rem), 2))
        return grouped[::-1] + "," + last_three
    except Exception:
        return amount_str

def format_indian_currency_with_decimal(value):
    try:
        f_val = float(str(value).replace(',', ''))
        parts = f"{f_val:.2f}".split('.')
        integer_part = parts[0]
        decimal_part = parts[1]
        
        if len(integer_part) <= 3:
            return integer_part + "." + decimal_part
            
        last_three = integer_part[-3:]
        remaining = integer_part[:-3]
        reversed_rem = remaining[::-1]
        grouped = ",".join(reversed_rem[i:i+2] for i in range(0, len(reversed_rem), 2))
        return grouped[::-1] + "," + last_three + "." + decimal_part
    except Exception:
        return str(value)

def increment_date(current_date, period_type):
    import calendar
    from datetime import timedelta
    pt = str(period_type).lower().strip()
    if pt == 'daily':
        return current_date + timedelta(days=1)
    elif pt == 'weekly':
        return current_date + timedelta(days=7)
    elif pt == 'biweekly':
        return current_date + timedelta(days=14)
    elif pt == 'bimonthly':
        return current_date + timedelta(days=15)
    elif pt == 'monthly':
        month = current_date.month
        year = current_date.year
        day = current_date.day
        
        month += 1
        if month > 12:
            month = 1
            year += 1
            
        last_day = calendar.monthrange(year, month)[1]
        if day > last_day:
            day = last_day
            
        return datetime(year, month, day)
    else:
        return current_date + timedelta(days=30)

def populate_table_rows(table, data_list, include_total=False, total_amount=0):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    hdr_row = table.rows[0]
    num_cols = len(hdr_row.cells)
    
    target_rows_count = 1 + len(data_list)
    if include_total:
        target_rows_count += 1
        
    for idx, data_row in enumerate(data_list):
        if idx + 1 < len(table.rows):
            row = table.rows[idx + 1]
        else:
            row = table.add_row()
            
        for col_idx in range(min(num_cols, len(data_row))):
            cell = row.cells[col_idx]
            cell.text = str(data_row[col_idx])
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
    if include_total:
        if len(table.rows) < target_rows_count:
            total_row = table.add_row()
        else:
            total_row = table.rows[target_rows_count - 1]
        total_row.cells[0].text = ''
        total_row.cells[1].text = 'TOTAL'
        total_row.cells[2].text = ''
        total_row.cells[3].text = format_indian_currency_with_decimal(total_amount)
        
        for cell in total_row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    while len(table.rows) > target_rows_count:
        row_to_remove = table.rows[-1]
        table._tbl.remove(row_to_remove._tr)

def merge_adjacent_runs(paragraph):
    if not paragraph.runs:
        return
    runs = paragraph.runs
    new_runs_data = []
    
    current_text = runs[0].text
    current_bold = runs[0].bold
    current_italic = runs[0].italic
    current_font_name = runs[0].font.name
    current_font_size = runs[0].font.size
    
    for run in runs[1:]:
        bold = run.bold
        italic = run.italic
        font_name = run.font.name
        font_size = run.font.size
        
        if (bold == current_bold and 
            italic == current_italic and 
            font_name == current_font_name and 
            font_size == current_font_size):
            current_text += run.text
        else:
            new_runs_data.append((current_text, current_bold, current_italic, current_font_name, current_font_size))
            current_text = run.text
            current_bold = bold
            current_italic = italic
            current_font_name = font_name
            current_font_size = font_size
            
    new_runs_data.append((current_text, current_bold, current_italic, current_font_name, current_font_size))
    
    p_element = paragraph._p
    for r in list(runs):
        p_element.remove(r._r)
        
    for text, bold, italic, font_name, font_size in new_runs_data:
        new_run = paragraph.add_run(text)
        new_run.bold = bold
        new_run.italic = italic
        if font_name:
            new_run.font.name = font_name
        if font_size:
            new_run.font.size = font_size

def apply_interest_underlining(paragraph, interest_value):
    import re
    for run in paragraph.runs:
        text = run.text
        if "18.00" in text and "(Flat)" in text:
            match = re.search(r"^(.*?)(18\.00)(\s*)(\(Flat\))(.*)$", text)
            if match:
                prefix, old_int, spaces, flat, suffix = match.groups()
                
                run.text = prefix
                
                r_int = paragraph.add_run(interest_value)
                r_int.underline = True
                r_int.bold = run.bold
                r_int.italic = run.italic
                r_int.font.name = run.font.name
                r_int.font.size = run.font.size
                
                r_spaces = paragraph.add_run(spaces)
                r_spaces.bold = run.bold
                r_spaces.italic = run.italic
                r_spaces.font.name = run.font.name
                r_spaces.font.size = run.font.size
                
                r_flat = paragraph.add_run(flat)
                r_flat.underline = True
                r_flat.bold = run.bold
                r_flat.italic = run.italic
                r_flat.font.name = run.font.name
                r_flat.font.size = run.font.size
                
                r_suf = paragraph.add_run(suffix)
                r_suf.bold = run.bold
                r_suf.italic = run.italic
                r_suf.font.name = run.font.name
                r_suf.font.size = run.font.size
                
                p_element = paragraph._p
                new_runs = [r_int, r_spaces, r_flat, r_suf]
                for r_new in new_runs:
                    p_element.remove(r_new._r)
                
                current_r_element = run._r
                for r_new in new_runs:
                    current_r_element.addnext(r_new._r)
                    current_r_element = r_new._r
                break

def format_company_address(addr_str):
    import re
    if not addr_str:
        return ""
    addr_str = re.sub(r'\s+', ' ', str(addr_str)).strip()
    parts = [p.strip() for p in addr_str.split(',') if p.strip()]
    if len(parts) >= 4:
        if "No." in parts[0] or parts[0].isdigit() or len(parts[0]) <= 10:
            line1 = f"{parts[0]}, {parts[1]},"
            start_idx = 2
        else:
            line1 = f"{parts[0]},"
            start_idx = 1
            
        lines = [line1]
        for p in parts[start_idx:-1]:
            lines.append(f"{p},")
        last_part = parts[-1]
        if not last_part.endswith('.'):
            last_part += '.'
        lines.append(last_part)
        return "\n".join(lines)
    return addr_str


def format_address_multiline(addr_str):
    """
    Cleans and structures an address string into a beautifully balanced,
    comma-separated multi-line address. Preserves user-provided newlines
    if they are present.
    """
    import re
    if not addr_str:
        return ""
        
    addr_str = str(addr_str).strip()
    
    # If the user already formatted the address with newlines, clean up each line
    if '\n' in addr_str:
        lines = [line.strip() for line in addr_str.split('\n') if line.strip()]
        cleaned_lines = []
        for line in lines:
            l = re.sub(r'^\s*,\s*', '', line).strip()
            l = re.sub(r'\s*,\s*$', '', l).strip()
            l = re.sub(r'\s+', ' ', l)
            if l:
                cleaned_lines.append(l)
        return "\n".join(cleaned_lines)
        
    # Otherwise, split by commas and group them beautifully
    parts = [p.strip() for p in addr_str.split(',') if p.strip()]
    parts = [re.sub(r'\s+', ' ', p) for p in parts]
    
    if not parts:
        return ""
        
    n = len(parts)
    if n <= 1:
        return parts[0]
    elif n == 2:
        return f"{parts[0]},\n{parts[1]}"
    elif n == 3:
        return f"{parts[0]},\n{parts[1]},\n{parts[2]}"
    elif n == 4:
        return f"{parts[0]}, {parts[1]},\n{parts[2]},\n{parts[3]}"
    else:
        # Group first two parts together, then third, then remaining
        line1 = f"{parts[0]}, {parts[1]}"
        line2 = parts[2]
        line3 = ", ".join(parts[3:])
        return f"{line1},\n{line2},\n{line3}"

def replace_address_in_paragraph(paragraph, ph, address_val):
    """
    Performs placeholder substitution in python-docx paragraphs while injecting
    line breaks, handling double "No." prefixing, and applying intermediate commas.
    """
    import re
    if ph not in paragraph.text:
        return False
        
    lines = [l.strip() for l in address_val.split('\n') if l.strip()]
    if not lines:
        return False
        
    first_line = lines[0]
    starts_with_number_or_no = False
    if re.search(r'^\d', first_line) or re.search(r'^no\.?\s*\d', first_line, re.IGNORECASE):
        starts_with_number_or_no = True
        
    # If the first line doesn't start with a number/No, strip the hardcoded 'No. ' prefix from paragraph runs
    if not starts_with_number_or_no:
        for run in paragraph.runs:
            for no_pattern in ['No. ', 'No.', 'no. ', 'no.', 'No.- ', 'No.-']:
                if no_pattern in run.text:
                    run.text = run.text.replace(no_pattern, '')
    else:
        # Address starts with No, strip leading No from address first line to avoid duplicating with template
        if first_line.upper().startswith('NO.'):
            first_line = re.sub(r'^no\.?\s*', '', first_line, flags=re.IGNORECASE).strip()
            lines[0] = first_line
        elif first_line.upper().startswith('NO '):
            first_line = re.sub(r'^no\s+', '', first_line, flags=re.IGNORECASE).strip()
            lines[0] = first_line

    # Replace the placeholder in the target run and inject multiline layout
    for run in list(paragraph.runs):
        if ph in run.text:
            # Clear double commas/dots in placeholder
            run.text = run.text.replace(f"{ph},", ph).replace(f"{ph}.", ph)
            parts = run.text.split(ph)
            run.text = parts[0]
            
            current_run = run
            for i, line in enumerate(lines):
                if i > 0:
                    current_run = paragraph.add_run()
                    current_run.add_break()
                    current_run.bold = run.bold
                    current_run.italic = run.italic
                    current_run.font.name = run.font.name
                    current_run.font.size = run.font.size
                    
                display_line = line
                if i < len(lines) - 1:
                    if not display_line.endswith(','):
                        display_line += ','
                else:
                    if not display_line.endswith('.'):
                        display_line += '.'
                        
                current_run = paragraph.add_run(display_line)
                current_run.bold = run.bold
                current_run.italic = run.italic
                current_run.font.name = run.font.name
                current_run.font.size = run.font.size
                
            if len(parts) > 1 and parts[1]:
                r_after = paragraph.add_run(parts[1])
                r_after.bold = run.bold
                r_after.italic = run.italic
                r_after.font.name = run.font.name
                r_after.font.size = run.font.size
            return True
    return False

def get_lender_details_from_excel(lender_name):
    """
    Look up company address and PAN from company_address_data.xlsx using the lender name.
    """
    import os
    import openpyxl
    import re
    
    excel_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), 'company_address_data.xlsx')
    default_details = {
        "address": "Bengaluru",
        "pan": "AATFJ7144B" # default fallback
    }
    if not os.path.exists(excel_path):
        return default_details
        
    try:
        wb = openpyxl.load_workbook(excel_path, read_only=True)
        sheet = wb.active
        
        target = lender_name.strip().upper()
        
        def clean_name(n):
            if not n:
                return ""
            n = str(n).upper().strip()
            # Remove M/S. or M/S or MR. or MRS.
            n = re.sub(r'^(M/S\.?|MR\.?|MRS\.?)\s+', '', n)
            return n.strip()
            
        clean_target = clean_name(target)
        
        for row in sheet.iter_rows(min_row=2, values_only=True):
            if not row or len(row) < 3:
                continue
            name_val = row[0]
            pan_val = row[1]
            addr_val = row[2]
            if name_val:
                clean_row_name = clean_name(name_val)
                if clean_row_name == clean_target:
                    address = str(addr_val).replace('\xa0', ' ').strip() if addr_val else "Bengaluru"
                    address = re.sub(r'\s+', ' ', address)
                    # Format to multi-line layout
                    address = format_company_address(address)
                    pan = str(pan_val).strip().upper() if pan_val else ""
                    return {
                        "address": address,
                        "pan": pan
                    }
                    
        return default_details
    except Exception as e:
        print(f"Error loading lender details from excel: {e}")
        return default_details

def fill_promissory_note_docx(form_data, joinees_list, template_path, output_path):
    import docx
    from datetime import datetime
    import re
    import os
    
    # 1. Load document
    doc = docx.Document(template_path)
    
    # Merge all paragraph runs globally first to simplify placeholders
    for p in doc.paragraphs:
        merge_adjacent_runs(p)
        
    # 2. Extract and format values (with Indian Currency style, NO decimals!)
    loan_amount_val = form_data.get('loanAmount', '')
    loan_amount_formatted = format_indian_currency(loan_amount_val)
        
    loan_date_val = form_data.get('loanDate', '')
    try:
        dt = datetime.strptime(loan_date_val, "%Y-%m-%d")
        loan_date_formatted = dt.strftime("%d-%m-%Y")
    except Exception:
        loan_date_formatted = loan_date_val
        
    amount_in_words = num_to_words(loan_amount_val)
    if amount_in_words.endswith(" Rupees Only"):
        amount_in_words = amount_in_words.replace(" Rupees Only", "")
        
    # Extract proprietor components
    proprietor_name = form_data.get('proprietorName', '')
    # Strip any Mr./Mrs./Smt./Sri title prefixes if already included
    clean_prop_name = proprietor_name
    for prefix in ['Mr.', 'Mrs.', 'Ms.', 'Sri.', 'Smt.', 'Sri', 'Smt', 'Mr', 'Mrs']:
        if clean_prop_name.upper().startswith(prefix.upper()):
            clean_prop_name = clean_prop_name[len(prefix):].strip()
            break
            
    proprietor_relation = 'D/o' if str(proprietor_name).strip().lower().startswith(('mrs', 'smt')) else 'S/o'
    proprietor_title = 'Sri' if proprietor_relation == 'S/o' else 'Smt'
    
    # 3. Handle Dynamic Joinee Substitution in Paragraph 5
    p5 = None
    for p in doc.paragraphs:
        if "ON DEMAND" in p.text:
            p5 = p
            break
    if p5 is None:
        p5 = doc.paragraphs[5] if len(doc.paragraphs) > 5 else doc.paragraphs[-1]
    
    # Underline the interest rate and (Flat) text
    interest_val = form_data.get('interest', '18')
    try:
        interest_formatted = f"{float(interest_val):.2f}"
    except Exception:
        interest_formatted = str(interest_val)
    apply_interest_underlining(p5, interest_formatted)
    
    # Clean title placeholders in merged runs of Paragraph 5 if title is already included in names
    company_name = form_data.get('companyName', '')
    if str(company_name).strip().lower().startswith(('m/s', 'm/s.')):
        for run in p5.runs:
            if "{{client_company_title}}" in run.text:
                run.text = run.text.replace("{{client_company_title}}.", "").replace("{{client_company_title}}", "")
        
    if str(proprietor_name).strip().lower().startswith(('mr', 'ms', 'mrs', 'sri', 'smt', 'dr')):
        for run in p5.runs:
            if "{{name_title}}" in run.text:
                run.text = run.text.replace("{{name_title}}.", "").replace("{{name_title}}", "")
        
    lender_name = form_data.get('lenderName', '')
    if str(lender_name).strip().lower().startswith(('m/s', 'm/s.')):
        for run in p5.runs:
            if "{{company_title}}" in run.text:
                run.text = run.text.replace("{{company_title}}.", "").replace("{{company_title}}", "")
        
    if not joinees_list:
        for run in p5.runs:
            if " and " in run.text:
                run.text = run.text.replace(" and ", "")
            if "{{joinee_title}}" in run.text or "{{joinee_name}}" in run.text:
                run.text = ""
            if "{{joinee_name_relation}}" in run.text or "{{joinee_father_name}}" in run.text or "{{joinee_address}}" in run.text:
                run.text = run.text.replace(" {{joinee_name_relation}}. {{joinee_father_name}}, {{joinee_address}},", "")
    elif len(joinees_list) == 1:
        j_name = joinees_list[0].get('name', '')
        if str(j_name).strip().lower().startswith(('mr', 'ms', 'mrs', 'sri', 'smt', 'dr')):
            for run in p5.runs:
                if "{{joinee_title}}" in run.text:
                    run.text = run.text.replace("{{joinee_title}}.", "").replace("{{joinee_title}}", "")
    else:
        # Multiple joinees
        joinee_str_parts = []
        for idx, j in enumerate(joinees_list):
            j_name = j.get('name', '')
            j_rel = "D/o" if str(j_name).strip().lower().startswith(('mrs', 'smt')) else "S/o"
            j_title = ""
            if not str(j_name).strip().lower().startswith(('mr', 'ms', 'mrs', 'sri', 'smt', 'dr')):
                j_title = "Sri. " if j_rel == "S/o" else "Smt. "
            j_father = j.get('father', '')
            j_addr = j.get('address', '')
            joinee_str_parts.append(f"({idx + 2}) {j_title}{j_name}, {j_rel}. {j_father}, {j_addr}")
        
        multi_joinee_full_string = " and " + " and ".join(joinee_str_parts) + ", "
        
        joinee_replaced = False
        for run in p5.runs:
            if " and " in run.text:
                run.text = run.text.replace(" and ", "")
            if "{{joinee_title}}" in run.text or "{{joinee_name}}" in run.text or "{{joinee_name_relation}}" in run.text:
                if not joinee_replaced:
                    run.text = multi_joinee_full_string
                    joinee_replaced = True
                else:
                    run.text = ""
        
    # Look up lender address and PAN from Excel
    lender_name_val = form_data.get('lenderName', 'Jubilant Enterprises Private Limited')
    lender_details = get_lender_details_from_excel(lender_name_val)
    
    # Format client address to a single continuous line and strip duplicate leading "No."
    client_addr = form_data.get('companyAddress', '')
    client_addr_clean = re.sub(r'^no\.?\s*', '', client_addr.replace('\n', ' '), flags=re.IGNORECASE).strip()
    client_addr_clean = re.sub(r'\s+', ' ', client_addr_clean)
    
    # Format lender address to a single continuous line and strip duplicate leading "No."
    lender_addr = lender_details['address']
    lender_addr_clean = re.sub(r'^no\.?\s*', '', lender_addr.replace('\n', ' '), flags=re.IGNORECASE).strip()
    lender_addr_clean = re.sub(r'\s+', ' ', lender_addr_clean)
    
    # 5. Prepare mapping dictionary
    mapping = {
        '{{loan_amount}}': loan_amount_formatted,
        '{{place_name}}': form_data.get('place', ''),
        '{{loan_date}}': loan_date_formatted,
        '{{client_company_title}}': 'M/s',
        '{{client_company_name}}': form_data.get('companyName', ''),
        '{{client_company_address}}': client_addr_clean,
        '{{name_title}}': proprietor_title,
        '{{name}}': clean_prop_name,
        '{{name_relation}}': proprietor_relation,
        '{{father_name}}': form_data.get('fatherOfProprietor', ''),
        '{{company_title}}': 'M/s',
        '{{company_name}}': lender_name_val,
        '{{company_address}}': lender_addr_clean,
        '{{company_pan}}': lender_details['pan'],
        '{{amount_text}}': amount_in_words,
        '{{bank_name}}': form_data.get('bankName', ''),
        '{{account_number}}': form_data.get('accountNumber', ''),
        '{{branch_name}}': form_data.get('branch', ''),
        '{{ifsc_code}}': form_data.get('ifsc', ''),
    }
    
    if len(joinees_list) == 1:
        j0_name = joinees_list[0].get('name', '')
        j0_relation = 'D/o' if str(j0_name).strip().lower().startswith(('mrs', 'smt')) else 'S/o'
        j0_title = 'Sri' if j0_relation == 'S/o' else 'Smt'
        mapping.update({
            '{{joinee_title}}': j0_title,
            '{{joinee_name}}': j0_name,
            '{{joinee_name_relation}}': j0_relation,
            '{{joinee_father_name}}': joinees_list[0].get('father', ''),
            '{{joinee_address}}': joinees_list[0].get('address', '')
        })
        
    # 5. Do replacement across all paragraphs and runs
    for paragraph in doc.paragraphs:
        for ph, val in mapping.items():
            for run in paragraph.runs:
                if ph in run.text:
                    run.text = run.text.replace(ph, str(val))
                    
    # Underline all runs in the title paragraph (Paragraph 0)
    if len(doc.paragraphs) > 0:
        for run in doc.paragraphs[0].runs:
            run.underline = True
        
    # Underline the top-left amount run in Paragraph 2
    if len(doc.paragraphs) > 2:
        for run in doc.paragraphs[2].runs:
            if "Rs." in run.text:
                run.underline = True
                        
    # 6. Enforce font family 'Calibri' and font size '14pt' globally
    from docx.shared import Pt
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(14)
                        
    # 7. Save the filled document
    doc.save(output_path)


def set_table_borders(table):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
        
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = tblBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tblBorders.append(border)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')

def create_and_populate_table(doc, target_p, installments, include_total, total_amount=0):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt, Inches, Cm
    
    table = doc.add_table(rows=1, cols=4)
    # Apply standard borders via direct XML manipulation to guarantee borders show up in all templates
    try:
        set_table_borders(table)
    except Exception:
        pass
        
    try:
        table.style = 'Table Grid'
    except Exception:
        # Fallback to default if 'Table Grid' style doesn't exist in the document's styles collection
        pass
        
    # Enforce Auto-fit and define beautiful, standard column widths spanning 6.5" printable area
    table.allow_autofit = True
    col_widths = [Inches(0.8), Inches(1.8), Inches(1.9), Inches(2.0)]
    
    # Format Header Row
    headers = ['S.no', 'Date', 'CHEQUE NO.', 'AMOUNT']
    table.rows[0].height = Cm(1.0)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        cell.text = h
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Clear any document-level paragraph spacing defaults to guarantee exact vertical centering
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        for run in p.runs:
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            
    # Populate installments
    for idx, inst in enumerate(installments):
        row = table.add_row()
        row.height = Cm(0.6)
        for col_idx in range(4):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            cell.text = str(inst[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Clear any document-level paragraph spacing defaults to guarantee exact vertical centering
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                
    if include_total:
        total_row = table.add_row()
        total_row.height = Cm(0.6)
        total_row.cells[0].width = col_widths[0]
        total_row.cells[0].text = ''
        
        total_row.cells[1].width = col_widths[1]
        total_row.cells[1].text = 'TOTAL'
        
        total_row.cells[2].width = col_widths[2]
        total_row.cells[2].text = ''
        
        total_row.cells[3].width = col_widths[3]
        total_row.cells[3].text = format_indian_currency_with_decimal(total_amount)
        
        for col_idx in range(4):
            cell = total_row.cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Clear any document-level paragraph spacing defaults to guarantee exact vertical centering
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                
    # Move the table to be right after the target element (paragraph)
    target_p._element.addnext(table._tbl)
    return table

def fill_letterpad_docx(form_data, joinees_list, template_path, output_path):
    import docx
    from datetime import datetime
    import os
    
    # 1. Load document
    doc = docx.Document(template_path)
    
    # Merge all paragraph runs globally first to simplify placeholders
    for p in doc.paragraphs:
        merge_adjacent_runs(p)
        
    # 2. Extract and format values
    loan_amount_val = form_data.get('loanAmount', '')
    loan_amount_formatted = format_indian_currency(loan_amount_val)
        
    loan_date_val = form_data.get('loanDate', '')
    try:
        dt = datetime.strptime(loan_date_val, "%Y-%m-%d")
        loan_date_formatted = dt.strftime("%d-%m-%Y")
    except Exception:
        loan_date_formatted = loan_date_val
        
    amount_in_words = num_to_words(loan_amount_val)
    if amount_in_words.endswith(" Rupees Only"):
        amount_in_words = amount_in_words.replace(" Rupees Only", "")
        
    # Extract proprietor components
    proprietor_name = form_data.get('proprietorName', '')
    clean_prop_name = proprietor_name
    for prefix in ['Mr.', 'Mrs.', 'Ms.', 'Sri.', 'Smt.', 'Sri', 'Smt', 'Mr', 'Mrs']:
        if clean_prop_name.upper().startswith(prefix.upper()):
            clean_prop_name = clean_prop_name[len(prefix):].strip()
            break
            
    proprietor_title = form_data.get('proprietorTitle', 'Mr.')
    proprietor_relation = 'D/o' if proprietor_title.upper() in ['MRS.', 'MS.', 'MRS', 'MS'] else 'S/o'
    
    proprietor_details = f"{proprietor_title} {clean_prop_name}, {proprietor_relation}. {form_data.get('fatherOfProprietor', '')}"
    
    # Look up lender address and PAN from Excel
    lender_name_val = form_data.get('lenderName', 'Jubilant Enterprises Private Limited')
    lender_details = get_lender_details_from_excel(lender_name_val)
    
    # 4. Construct multi joinees text for Letterpad template
    multi_joinees_text = ""
    if joinees_list:
        joinee_str_parts = []
        for idx, j in enumerate(joinees_list):
            j_name = j.get('name', '')
            j_rel = "D/o" if str(j_name).strip().lower().startswith(('mrs', 'smt')) else "S/o"
            j_title = ""
            if not str(j_name).strip().lower().startswith(('mr', 'ms', 'mrs', 'sri', 'smt', 'dr')):
                j_title = "Mr. " if j_rel == "S/o" else "Mrs. "
            j_father = j.get('father', '')
            joinee_str_parts.append(f"({idx + 2}) {j_title}{j_name}, {j_rel}. {j_father}")
        multi_joinees_text = ", and " + ", and ".join(joinee_str_parts)

    # 5. Prepare mapping dictionary
    mapping = {
        '{{loan_amount}}': loan_amount_formatted,
        '{{place_name}}': form_data.get('place', ''),
        '{{loan_date}}': loan_date_formatted,
        '{{client_company_title}}': 'M/s',
        '{{client_company_name}}': form_data.get('companyName', ''),
        '{{proprietor_details}}': proprietor_details,
        '{{name_title}}': proprietor_title,
        '{{name}}': clean_prop_name,
        '{{name_relation}}': proprietor_relation,
        '{{father_name}}': form_data.get('fatherOfProprietor', ''),
        '{{company_name}}': lender_name_val,
        '{{company_pan}}': lender_details['pan'],
        '{{amount_text}}': amount_in_words,
        '{{bank_name}}': form_data.get('bankName', ''),
        '{{account_number}}': form_data.get('accountNumber', ''),
        '{{branch_name}}': form_data.get('branch', ''),
        '{{ifsc_code}}': form_data.get('ifsc', ''),
        '{{multi_joinees_text}}': multi_joinees_text,
    }
    
    company_address = form_data.get('companyAddress', '')
    company_address_formatted = format_address_multiline(company_address)
    
    lender_address = lender_details['address']
    lender_address_formatted = format_address_multiline(lender_address)
    
    # 5. Do replacement across all paragraphs and runs
    for paragraph in doc.paragraphs:
        replace_address_in_paragraph(paragraph, '{{client_company_address}}', company_address_formatted)
        replace_address_in_paragraph(paragraph, '{{company_address}}', lender_address_formatted)
        
        for ph, val in mapping.items():
            for run in paragraph.runs:
                if ph in run.text:
                    run.text = run.text.replace(ph, str(val))
                    
    # Generate and populate installment table dynamically
    emi_start_date_str = form_data.get('emiStartDate', '')
    if not emi_start_date_str:
        emi_start_date_str = form_data.get('loanDate', '')
    if not emi_start_date_str:
        emi_start_date_str = datetime.today().strftime("%Y-%m-%d")
        
    try:
        current_date = datetime.strptime(emi_start_date_str, "%Y-%m-%d")
    except Exception:
        try:
            current_date = datetime.strptime(emi_start_date_str, "%d-%m-%Y")
        except Exception:
            current_date = datetime.today()
            
    period_type = form_data.get('period', 'monthly')
    
    repayment_val = 0.0
    try:
        repayment_val = float(str(form_data.get('repayment', '0')).replace(',', ''))
    except Exception:
        pass
        
    no_of_periods = 1
    try:
        no_of_periods = int(form_data.get('noOfPeriod', '1'))
    except Exception:
        pass
        
    inst_amount = repayment_val / no_of_periods if no_of_periods > 0 else repayment_val
    
    installments = []
    current_inst_date = current_date
    for i in range(1, no_of_periods + 1):
        formatted_date = current_inst_date.strftime("%d-%m-%Y")
        formatted_amount = format_indian_currency_with_decimal(inst_amount)
        installments.append((str(i), formatted_date, '', formatted_amount))
        current_inst_date = increment_date(current_inst_date, period_type)
        
    # Search for the {{table}} placeholder paragraph
    target_paragraph = None
    for p in doc.paragraphs:
        if '{{table}}' in p.text:
            target_paragraph = p
            break
            
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    if target_paragraph:
        if no_of_periods <= 10:
            create_and_populate_table(doc, target_paragraph, installments, include_total=True, total_amount=repayment_val)
        else:
            # Create Table 0
            t0 = create_and_populate_table(doc, target_paragraph, installments[:10], include_total=False)
            
            # For multi-page, add the signature paragraph immediately after Table 0 (will sit at the bottom of Page 1)
            p_sig0 = doc.add_paragraph()
            p_sig0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_sig0.paragraph_format.space_before = Pt(200) # Push to the absolute bottom line of the first page
            p_sig0.paragraph_format.space_after = Pt(0)
            run_label0 = p_sig0.add_run("Signature of Borrower(s)")
            run_label0.font.name = 'Calibri'
            run_label0.font.size = Pt(12)
            run_label0.bold = True
            run_label0.italic = True
            run_label0.font.color.rgb = RGBColor(0, 0, 0)
            t0._element.addnext(p_sig0._element)
            
            # Create a manual page break between the two tables
            p_break = doc.add_page_break()
            
            # Position the page break after the signature paragraph
            p_sig0._element.addnext(p_break._element)
            
            # Create Table 1 directly after the page break
            create_and_populate_table(doc, p_break, installments[10:], include_total=True, total_amount=repayment_val)
            
        # Clean up the {{table}} placeholder paragraph from the document
        target_paragraph._element.getparent().remove(target_paragraph._element)
    else:
        # Fallback to pre-existing tables if {{table}} placeholder is not found
        if len(doc.tables) >= 2:
            table0 = doc.tables[0]
            table1 = doc.tables[1]
            
            if no_of_periods <= 10:
                populate_table_rows(table0, installments, include_total=True, total_amount=repayment_val)
                # Remove Table 1 safely from XML
                table1._element.getparent().remove(table1._element)
            else:
                populate_table_rows(table0, installments[:10], include_total=False)
                
                # Add signature paragraph after table0 (Page 1)
                p_sig0 = doc.add_paragraph()
                p_sig0.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_sig0.paragraph_format.space_before = Pt(200) # Push to the absolute bottom line of the first page
                p_sig0.paragraph_format.space_after = Pt(0)
                run_label0 = p_sig0.add_run("Signature of Borrower(s)")
                run_label0.font.name = 'Calibri'
                run_label0.font.size = Pt(12)
                run_label0.bold = True
                run_label0.italic = True
                run_label0.font.color.rgb = RGBColor(0, 0, 0)
                table0._element.addnext(p_sig0._element)
                
                # Add page break after the signature of Page 1
                p_break = doc.add_page_break()
                p_sig0._element.addnext(p_break._element)
                
                populate_table_rows(table1, installments[10:], include_total=True, total_amount=repayment_val)
        elif len(doc.tables) == 1:
            table0 = doc.tables[0]
            populate_table_rows(table0, installments, include_total=True, total_amount=repayment_val)
            
    # Add 'Signature of Borrower(s)' at the very end of the document (below "Yours faithfully,")
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Mathematically calculate remaining space on Page 2 based on the number of rows in Table 1
    # 450pt is the approximate printable height. Each row takes about 17pt (0.6cm).
    # We cap the spacing between Pt(36) and Pt(200) to ensure it is always perfectly bottom-aligned and never overflows.
    if no_of_periods > 10:
        rows_page2 = no_of_periods - 10
        calculated_spacing = 450 - (rows_page2 * 17)
        last_page_spacing = Pt(min(200, max(36, calculated_spacing)))
    else:
        last_page_spacing = Pt(200)
        
    p_sig.paragraph_format.space_before = last_page_spacing
    p_sig.paragraph_format.space_after = Pt(0)
    run_label = p_sig.add_run("Signature of Borrower(s)")
    run_label.font.name = 'Calibri'
    run_label.font.size = Pt(12)
    run_label.bold = True
    run_label.italic = True
    run_label.font.color.rgb = RGBColor(0, 0, 0)
                    
    # 6. Enforce font family 'Calibri' and font size '12pt' globally for Letterpad
    from docx.shared import Pt
    from docx.oxml.ns import qn
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), 'Calibri')
            rFonts.set(qn('w:hAnsi'), 'Calibri')
            run.font.size = Pt(12)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        rPr = run._r.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:ascii'), 'Calibri')
                        rFonts.set(qn('w:hAnsi'), 'Calibri')
                        run.font.size = Pt(12)
                        
    # 7. Save the filled document
    doc.save(output_path)

def format_guarantor_address(addr_str):
    import re
    # 1. Strip newlines and spaces
    addr = addr_str.replace('\n', ' ').strip()
    addr = re.sub(r'\s+', ' ', addr)
    addr = re.sub(r'^no\.?\s*', '', addr, flags=re.IGNORECASE)
    
    # 2. Find pincode (6-digit number)
    pin_match = re.search(r'\b\d{6}\b', addr)
    pincode = pin_match.group(0) if pin_match else ""
    
    # 3. Remove pincode from the address string to avoid duplicate matching
    if pincode:
        addr = addr.replace(pincode, "")
        
    # 4. Identify state (Kerala or Tamil Nadu or standard Indian states)
    state_names = ["Kerala", "Tamil Nadu", "Tamilnadu", "Karnataka", "Andhra Pradesh", "Telangana"]
    state = ""
    for s in state_names:
        match = re.search(r'\b' + re.escape(s) + r'\b', addr, flags=re.IGNORECASE)
        if match:
            state = match.group(0)
            addr = addr.replace(match.group(0), "")
            break
            
    # Clean up state representation
    if state.lower() == 'tamilnadu':
        state = 'Tamil Nadu'
    elif state:
        state = state.title()
        
    # 5. Split remaining address by commas or spaces and clean up empty elements
    addr = re.sub(r'[\s,\-\.]+$', '', addr)
    addr = re.sub(r'^[\s,\-\.]+', '', addr)
    
    parts = [p.strip() for p in addr.split(',') if p.strip()]
    parts = [re.sub(r'[\-\s]+$', '', p).strip() for p in parts]
    parts = [p for p in parts if p]
    
    if len(parts) >= 1:
        district = parts[-1]
        other_parts = parts[:-1]
    else:
        district = ""
        other_parts = []
        
    # 6. Reconstruct: [Other Parts], [District] - [Pincode], [State].
    reconstructed_addr = ", ".join(other_parts)
    if reconstructed_addr:
        reconstructed_addr += ", "
        
    if district:
        reconstructed_addr += district
        
    if pincode:
        if district:
            reconstructed_addr += f" - {pincode}"
        else:
            reconstructed_addr += f"{pincode}"
            
    if state:
        if district or pincode:
            reconstructed_addr += f", {state}"
        else:
            reconstructed_addr += f"{state}"
            
    if not reconstructed_addr.endswith('.'):
        reconstructed_addr += "."
        
    return reconstructed_addr

def fill_ltrl_docx(form_data, joinees_list, template_path, output_path):
    import docx
    from datetime import datetime
    import re
    import os
    from docx.shared import Pt
    from docx.oxml.ns import qn
    
    doc = docx.Document(template_path)
    
    # Extract values
    loan_amount_val = form_data.get('loanAmount', '')
    if '.' in str(loan_amount_val):
        loan_amount_val = str(loan_amount_val).split('.')[0]
        
    loan_amount_formatted = format_indian_currency(loan_amount_val)
    amount_in_words = num_to_words(loan_amount_val)
    if amount_in_words.endswith(" Rupees Only"):
        amount_in_words = amount_in_words.replace(" Rupees Only", "")
        
    loan_date_val = form_data.get('loanDate', '')
    try:
        dt = datetime.strptime(loan_date_val, "%Y-%m-%d")
        loan_date_formatted = dt.strftime("%d-%m-%Y")
    except Exception:
        loan_date_formatted = loan_date_val

    # Company name and formatted address
    company_name = form_data.get('companyName', '')
    company_address = form_data.get('companyAddress', '')
    company_address_formatted = format_address_multiline(company_address)

    # Proprietor
    proprietor_name = form_data.get('proprietorName', '')
    prop_relation = 'D/o' if str(proprietor_name).strip().lower().startswith(('mrs', 'smt')) else 'S/o'
    prop_title = 'Mr.' if prop_relation == 'S/o' else 'Mrs.'
    
    # Strip prefixes from clean prop name
    clean_prop_name = proprietor_name
    for prefix in ['Mr.', 'Mrs.', 'Ms.', 'Sri.', 'Smt.', 'Sri', 'Smt', 'Mr', 'Mrs']:
        if clean_prop_name.upper().startswith(prefix.upper()):
            clean_prop_name = clean_prop_name[len(prefix):].strip()
            break

    prop_father = form_data.get('fatherOfProprietor', '')
    prop_pan = form_data.get('proprietorPan', '').upper()

    # Lender details from Excel
    lender_name_val = form_data.get('lenderName', 'JUBILANT CAPITAL')
    lender_details = get_lender_details_from_excel(lender_name_val)
    lender_pan = lender_details.get('pan', '')
    lender_address = lender_details.get('address', '')
    
    # Strip lender name and address clean
    lender_name_clean = re.sub(r'^(M/S\.?)\s+', '', lender_name_val, flags=re.IGNORECASE).strip()
    lender_address_clean = re.sub(r'^no\.?\s*', '', lender_address.replace('\n', ' '), flags=re.IGNORECASE).strip()
    lender_address_clean = re.sub(r'\s+', ' ', lender_address_clean)

    # 1. Update Date
    doc.paragraphs[0].text = f"From\t\t\t\t\t\t\t\t\t\tDate: {loan_date_formatted}"
    
    # 2. Update Borrower Firm Name & Address (Make company name bold)
    p1 = doc.paragraphs[1]
    p1.text = ""
    p1.add_run("M/s. ")
    r_comp = p1.add_run(company_name)
    r_comp.bold = True
    p1.add_run(", ")
    
    p2 = doc.paragraphs[2]
    p2.text = ""
    
    lines = [l.strip() for l in company_address_formatted.split('\n') if l.strip()]
    first_line = lines[0] if lines else ""
    starts_with_number_or_no = False
    if first_line:
        if re.search(r'^\d', first_line) or re.search(r'^no\.?\s*\d', first_line, re.IGNORECASE):
            starts_with_number_or_no = True
            
    if starts_with_number_or_no and first_line:
        if first_line.upper().startswith('NO.'):
            first_line = re.sub(r'^no\.?\s*', '', first_line, flags=re.IGNORECASE).strip()
            lines[0] = first_line
        elif first_line.upper().startswith('NO '):
            first_line = re.sub(r'^no\s+', '', first_line, flags=re.IGNORECASE).strip()
            lines[0] = first_line
            
    if starts_with_number_or_no:
        p2.add_run("No. ")
        
    for i, line in enumerate(lines):
        if i > 0:
            p2.add_run().add_break()
        display_line = line
        if i < len(lines) - 1:
            if not display_line.endswith(','):
                display_line += ','
        else:
            if not display_line.endswith(','):
                display_line += ','
        p2.add_run(display_line)
    
    # 3. Update Proprietor Name, Relation, PAN, Father (Make proprietor name bold)
    p5 = doc.paragraphs[5]
    p5.text = ""
    p5.add_run(f"(1) {prop_title} ")
    r_prop = p5.add_run(clean_prop_name)
    r_prop.bold = True
    p5.add_run(f", (PAN No. {prop_pan}), (Proprietor),")
    
    # 4. Handle Joinees (Make guarantor name bold)
    if not joinees_list:
        doc.paragraphs[6].text = f"{prop_relation}. {prop_father}."
        doc.paragraphs[8].text = ""
        doc.paragraphs[9].text = ""
        doc.paragraphs[10].text = ""
        if len(doc.paragraphs) > 48:
            doc.paragraphs[48].text = "" # Clear signature 2
    else:
        doc.paragraphs[6].text = f"{prop_relation}. {prop_father}. \t\t\t\t\tand"
        j = joinees_list[0]
        j_name = j.get('name', '')
        # Strip any prefix titles from guarantor name to prevent double titles
        clean_j_name = j_name
        for prefix in ['Mr.', 'Mrs.', 'Ms.', 'Sri.', 'Smt.', 'Sri', 'Smt', 'Mr', 'Mrs']:
            if clean_j_name.upper().startswith(prefix.upper()):
                clean_j_name = clean_j_name[len(prefix):].strip()
                break
                
        j_title = j.get('title', 'Mr.')
        if j_title and not j_title.endswith('.'):
            j_title += '.'
        j_rel = 'D/o' if j_title.strip().lower() in ['mrs.', 'ms.', 'smt.'] else 'S/o'
        j_father = j.get('father', '')
        j_pan = j.get('pan', '').upper()
        j_addr = j.get('address', '')
        j_addr_clean = format_guarantor_address(j_addr)
        
        p8 = doc.paragraphs[8]
        p8.text = ""
        p8.add_run(f"(2) {j_title} ")
        r_guar = p8.add_run(clean_j_name)
        r_guar.bold = True
        p8.add_run(f", (Pan. No. {j_pan})")
        
        doc.paragraphs[9].text = f"{j_rel}. {j_father}, "
        doc.paragraphs[10].text = f"{j_addr_clean}."
        if len(doc.paragraphs) > 48:
            doc.paragraphs[48].text = "2." # Restore signature 2

    # 5. Update Lender details (Make lender name bold)
    p13 = doc.paragraphs[13]
    p13.text = ""
    p13.add_run("M/s. ")
    r_lend = p13.add_run(lender_name_clean)
    r_lend.bold = True
    p13.add_run(f", (PAN No. {lender_pan}),")
    
    # Split lender address into lines to fit nicely
    lender_address_lines = [p.strip() for p in lender_address.split(',') if p.strip()]
    
    # Fill in paragraphs 14-18
    for idx, line in enumerate(lender_address_lines[:5]):
        doc.paragraphs[14 + idx].text = line + ("," if idx < len(lender_address_lines) - 1 else ".")
    for idx in range(len(lender_address_lines), 5):
        doc.paragraphs[14 + idx].text = ""
        
    # 6. Update main body paragraph (Paragraph 23) (Make company, proprietor, guarantor names bold)
    p23 = doc.paragraphs[23]
    p23.text = "" # Clear template text to build run segments
    
    p23.add_run("I / We (1) ")
    r_prop23 = p23.add_run(f"{prop_title} {clean_prop_name}")
    r_prop23.bold = True
    p23.add_run(f", {prop_relation}. {prop_father} proprietor for and behalf of M/s. ")
    
    r_comp23 = p23.add_run(company_name)
    r_comp23.bold = True
    
    if joinees_list:
        j = joinees_list[0]
        j_name = j.get('name', '')
        # Strip prefixes here too for body request co-applicant text
        clean_j_name23 = j_name
        for prefix in ['Mr.', 'Mrs.', 'Ms.', 'Sri.', 'Smt.', 'Sri', 'Smt', 'Mr', 'Mrs']:
            if clean_j_name23.upper().startswith(prefix.upper()):
                clean_j_name23 = clean_j_name23[len(prefix):].strip()
                break
                
        j_title = j.get('title', 'Mr.')
        if j_title and not j_title.endswith('.'):
            j_title += '.'
        j_rel = 'D/o' if j_title.strip().lower() in ['mrs.', 'ms.', 'smt.'] else 'S/o'
        j_father = j.get('father', '')
        p23.add_run(" and (2) ")
        r_guar23 = p23.add_run(f"{j_title} {clean_j_name23}")
        r_guar23.bold = True
        p23.add_run(f", {j_rel}. {j_father}")
        
    p23.add_run(f" kindly request you to release the loan amount Rs. {loan_amount_formatted}/- (Rupees {amount_in_words} Only) requested by me / us vide application enclosed herewith and payable to me / us, as ")
    
    run_strike = p23.add_run("Account Payee / & CO. / ______________ Cheque / Demand Draft")
    run_strike.font.strike = True
    
    p23.add_run(" / ____RTGS____ payable to and at the following")
    
    # 7. Update bank details table labels (Make company name bold)
    p26 = doc.paragraphs[26]
    p26.text = ""
    p26.add_run("From Name: M/s. ")
    r_comp26 = p26.add_run(company_name)
    r_comp26.bold = True
    p26.add_run(", ")
    doc.paragraphs[27].text = f"Bank: {form_data.get('bankName', '')}, "
    doc.paragraphs[28].text = f"Account Number. - {form_data.get('accountNumber', '')}, "
    doc.paragraphs[29].text = f"Branch: {form_data.get('branch', '')}, "
    doc.paragraphs[30].text = f"IFSC Code: {form_data.get('ifsc', '')}."
    
    # 8. Force Calibri 12pt globally to look premium
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = 'Calibri'
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), 'Calibri')
            rFonts.set(qn('w:hAnsi'), 'Calibri')
            run.font.size = Pt(12)
            
    doc.save(output_path)

def generate_promissory_note_pdf(form_data, joinees_list, output_path):
    # Setup document geometry (A4, 0.75" margins)
    margin = 54
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin
    )
    
    styles = getSampleStyleSheet()
    
    # Custom Styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    meta_style = ParagraphStyle(
        'DocMeta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=11,
        leading=15,
        spaceAfter=15
    )
    
    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=15
    )
    
    # Format loan details
    loan_amount = form_data.get('loanAmount', '')
    loan_amount_formatted = format_indian_currency(loan_amount)
    
    loan_date_val = form_data.get('loanDate', '')
    try:
        dt = datetime.strptime(loan_date_val, "%Y-%m-%d")
        loan_date_formatted = dt.strftime("%d-%m-%Y")
    except Exception:
        loan_date_formatted = loan_date_val
        
    amount_in_words = num_to_words(loan_amount)
    if amount_in_words.endswith(" Rupees Only"):
        amount_in_words = amount_in_words.replace(" Rupees Only", "")
        
    interest_val = form_data.get('interest', '18')
    try:
        interest_formatted = f"{float(interest_val):.2f}"
    except Exception:
        interest_formatted = str(interest_val)
        
    place = form_data.get('place', '')
    company_name = form_data.get('companyName', '')
    company_address = form_data.get('companyAddress', '')
    proprietor_name = form_data.get('proprietorName', '')
    father_name = form_data.get('fatherOfProprietor', '')
    
    proprietor_relation = 'D/o' if str(proprietor_name).strip().lower().startswith(('mrs', 'smt')) else 'S/o'
    proprietor_title = 'Sri' if proprietor_relation == 'S/o' else 'Smt'
    
    # Construct joinees text
    joinees_text = ""
    if joinees_list:
        joinee_parts = []
        for idx, j in enumerate(joinees_list):
            j_name = j.get('name', '')
            j_rel = "D/o" if str(j_name).strip().lower().startswith(('mrs', 'smt')) else "S/o"
            j_title = ""
            if not str(j_name).strip().lower().startswith(('mr', 'ms', 'mrs', 'sri', 'smt', 'dr')):
                j_title = "Sri. " if j_rel == "S/o" else "Smt. "
            j_father = j.get('father', '')
            j_addr = j.get('address', '')
            joinee_parts.append(f"({idx + 2}) {j_title}{j_name}, {j_rel}. {j_father}, residing at {j_addr}")
        joinees_text = " and " + " and ".join(joinee_parts)
        
    lender_name = form_data.get('lenderName', 'Jubilant Enterprises Private Limited')
    
    # Story flowables
    story = []
    
    # 1. Main Title
    story.append(Paragraph("<u>PROMISSORY NOTE</u>", title_style))
    story.append(Spacer(1, 15))
    
    # 2. Header Metadata table (Amount, Place, Date)
    meta_data = [
        [Paragraph(f"<b>Rs. {loan_amount_formatted}/-</b>", meta_style), 
         Paragraph(f"<b>Place:</b> {place}<br/><b>Date:</b> {loan_date_formatted}", ParagraphStyle('MetaRight', parent=meta_style, alignment=TA_RIGHT))]
    ]
    meta_table = Table(meta_data, colWidths=[240, 240])
    meta_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
        ('TOPPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 20))
    
    # 3. Promissory text
    promissory_text = (
        f"ON DEMAND, I, <b>M/s. {company_name}</b>, having its principal place of business at "
        f"{company_address}, represented by its Proprietor <b>{proprietor_title}. {proprietor_name}</b>, "
        f"{proprietor_relation}. {father_name}{joinees_text}, promise to pay "
        f"<b>M/s. {lender_name}</b>, or order, the sum of "
        f"<b>Rs. {loan_amount_formatted}/-</b> (Rupees {amount_in_words} only) "
        f"together with interest at the rate of <b><u>{interest_formatted}% (Flat)</u></b> per annum "
        f"from the date of this note till the date of payment in full, for value received."
    )
    story.append(Paragraph(promissory_text, body_style))
    story.append(Spacer(1, 40))
    
    # 4. Signatures Table
    sig_data = [
        [Paragraph("<b>Co-applicant / Joinee</b><br/><br/>______________________", meta_style),
         Paragraph("<b>For M/s. " + company_name + "</b><br/><br/>______________________<br/>Proprietor", ParagraphStyle('SigRight', parent=meta_style, alignment=TA_RIGHT))]
    ]
    sig_table = Table(sig_data, colWidths=[240, 240])
    sig_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('LEFTPADDING', (0,0), (-1,-1), 0),
        ('RIGHTPADDING', (0,0), (-1,-1), 0),
    ]))
    story.append(KeepTogether(sig_table))
    
    # Build Document
    doc.build(story)

def fill_letter_of_undertaking_docx(form_data, joinees_list, template_path, output_path):
    import docx
    from datetime import datetime
    import re
    from docx.shared import Pt
    from docx.oxml.ns import qn
    
    # 1. Load document
    doc = docx.Document(template_path)
    
    # Merge runs globally to simplify search-and-replace
    for p in doc.paragraphs:
        merge_adjacent_runs(p)
        
    # 2. Extract values
    loan_amount_val = form_data.get('loanAmount', '')
    if '.' in str(loan_amount_val):
        loan_amount_val = str(loan_amount_val).split('.')[0]
    loan_amount_formatted = format_indian_currency(loan_amount_val)
    amount_in_words = num_to_words(loan_amount_val)
    if amount_in_words.endswith(" Rupees Only"):
        amount_in_words = amount_in_words.replace(" Rupees Only", "")
        
    loan_date_val = form_data.get('loanDate', '')
    try:
        dt = datetime.strptime(loan_date_val, "%Y-%m-%d")
        loan_date_formatted = dt.strftime("%d-%m-%Y")
    except Exception:
        loan_date_formatted = loan_date_val
        
    company_name = form_data.get('companyName', '').upper()
    company_address = form_data.get('companyAddress', '')
    company_address_formatted = format_address_multiline(company_address)
    
    proprietor_name = form_data.get('proprietorName', '').upper()
    prop_title = form_data.get('proprietorTitle', 'Mr.').strip()
    if prop_title and not prop_title.endswith('.'):
        prop_title += '.'
    prop_relation = 'D/o' if prop_title.strip().lower() in ['mrs.', 'ms.', 'smt.'] else 'S/o'
    prop_father = form_data.get('fatherOfProprietor', '').upper()
    prop_pan = form_data.get('proprietorPan', '').upper()
    
    # Strip prefixes from clean prop name
    clean_prop_name = proprietor_name
    for prefix in ['Mr.', 'Mrs.', 'Ms.', 'Sri.', 'Smt.', 'Sri', 'Smt', 'Mr', 'Mrs']:
        if clean_prop_name.upper().startswith(prefix.upper()):
            clean_prop_name = clean_prop_name[len(prefix):].strip()
            break
            
    # Lender details
    lender_name_val = form_data.get('lenderName', 'JUBILANT CAPITAL')
    lender_details = get_lender_details_from_excel(lender_name_val)
    lender_pan = lender_details.get('pan', '')
    lender_address = lender_details.get('address', '')
    
    lender_name_clean = re.sub(r'^(M/S\.?)\s*', '', lender_name_val, flags=re.IGNORECASE).strip()
    lender_address_formatted = format_address_multiline(lender_address)
    
    # Period term mapping
    period_val = form_data.get('period', '').lower()
    no_of_periods = form_data.get('noOfPeriod', '20')
    if 'month' in period_val:
        period_plural = 'months'
        period_term = 'monthly'
    elif 'week' in period_val:
        period_plural = 'weeks'
        period_term = 'weekly'
    elif 'day' in period_val:
        period_plural = 'days'
        period_term = 'daily'
    else:
        period_plural = 'periods'
        period_term = period_val
        
    mapping = {
        '{{loan_date}}': loan_date_formatted,
        '{{client_company_title}}': 'M/s',
        '{{client_company_name}}': company_name,
        '{{name_title}}': prop_title,
        '{{name}}': clean_prop_name,
        '{{proprietor_pan}}': prop_pan,
        '{{name_relation}}': prop_relation,
        '{{father_name}}': prop_father,
        '{{place_name}}': form_data.get('place', ''),
        '{{company_title}}': 'M/s',
        '{{company_name}}': lender_name_clean,
        '{{company_pan}}': lender_pan,
        '{{loan_amount}}': loan_amount_formatted,
        '{{amount_text}}': amount_in_words,
        '{{no_of_periods}}': no_of_periods,
        '{{period_plural}}': period_plural,
        '{{period_term}}': period_term,
    }
    
    if joinees_list:
        j = joinees_list[0]
        j_name = j.get('name', '')
        clean_j_name = j_name
        for prefix in ['Mr.', 'Mrs.', 'Ms.', 'Sri.', 'Smt.', 'Sri', 'Smt', 'Mr', 'Mrs']:
            if clean_j_name.upper().startswith(prefix.upper()):
                clean_j_name = clean_j_name[len(prefix):].strip()
                break
                
        j_title = j.get('title', 'Mr.')
        if j_title and not j_title.endswith('.'):
            j_title += '.'
        j_rel = 'D/o' if j_title.strip().lower() in ['mrs.', 'ms.', 'smt.'] else 'S/o'
        j_father = j.get('father', '').upper()
        j_pan = j.get('pan', '').upper()
        j_addr = j.get('address', '')
        j_addr_clean = format_guarantor_address(j_addr)
        
        mapping.update({
            '{{joinee_title}}': j_title,
            '{{joinee_name}}': clean_j_name,
            '{{joinee_pan}}': j_pan,
            '{{joinee_name_relation}}': j_rel,
            '{{joinee_father_name}}': j_father,
            '{{joinee_address}}': j_addr_clean,
        })
        
    # Search and replace in all paragraphs
    for paragraph in doc.paragraphs:
        # Perform address replacement
        replace_address_in_paragraph(paragraph, '{{client_company_address}}', company_address_formatted)
        replace_address_in_paragraph(paragraph, '{{company_address}}', lender_address_formatted)
        
        for ph, val in mapping.items():
            for run in paragraph.runs:
                if ph in run.text:
                    run.text = run.text.replace(ph, str(val))
                    
    # Format company name bold, proprietor name bold dynamically
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if company_name in run.text:
                run.bold = True
            if clean_prop_name in run.text:
                run.bold = True
            if joinees_list and clean_j_name in run.text:
                run.bold = True
                
    # 8. Align Signature block to page bottom (Signature of Borrower(s))
    p_sig_indices = []
    for idx, p in enumerate(doc.paragraphs):
        if "signature of" in p.text.lower() or "borrower(s)" in p.text.lower():
            p_sig_indices.append(idx)
            
    # Actually delete empty paragraphs in specified layout ranges to prevent page overflow
    for idx in range(min(52, len(doc.paragraphs) - 1), 42, -1):
        if idx < len(doc.paragraphs) and not doc.paragraphs[idx].text.strip():
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)
            
    for idx in range(min(32, len(doc.paragraphs) - 1), 24, -1):
        if idx < len(doc.paragraphs) and not doc.paragraphs[idx].text.strip():
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)
            
    # Find signature indices again after paragraph deletion
    p_sig_indices = []
    for idx, p in enumerate(doc.paragraphs):
        if "signature of" in p.text.lower() or "borrower(s)" in p.text.lower():
            p_sig_indices.append(idx)
        
    if p_sig_indices:
        first_sig_idx = p_sig_indices[0]
        if first_sig_idx < len(doc.paragraphs):
            p_sig1 = doc.paragraphs[first_sig_idx]
            if not joinees_list:
                p_sig1.paragraph_format.space_before = Pt(120)
            else:
                p_sig1.paragraph_format.space_before = Pt(90)
        
        if len(p_sig_indices) > 1:
            second_sig_idx = p_sig_indices[1]
            if second_sig_idx < len(doc.paragraphs):
                p_sig2 = doc.paragraphs[second_sig_idx]
                p_sig2.paragraph_format.space_before = Pt(100)
                
    # Delete any trailing empty paragraphs at the end of the document to guarantee no extra empty page is generated
    while len(doc.paragraphs) > 0 and not doc.paragraphs[-1].text.strip():
        p = doc.paragraphs[-1]
        p._element.getparent().remove(p._element)

    # 9. Force Calibri 12pt globally to look premium
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = 'Calibri'
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:ascii'), 'Calibri')
            rFonts.set(qn('w:hAnsi'), 'Calibri')
            run.font.size = Pt(12)
            
    doc.save(output_path)
