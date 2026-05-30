import docx
import sys

doc_path = r"d:\Projects\DocumentsTeam\backend\template\Proprietor\Letter of Undertaking.docx"
try:
    doc = docx.Document(doc_path)
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"{idx}: {p.text}")
except Exception as e:
    print("Error:", e)
